from docx import Document
import os
import glob
import argparse
import re

def split_by_chapters(input_file, output_dir, heading_style='Heading 1'):
    # 创建文件特定的输出目录
    file_name = os.path.splitext(os.path.basename(input_file))[0]
    file_output_dir = os.path.join(output_dir, file_name)
    if not os.path.exists(file_output_dir):
        os.makedirs(file_output_dir)
    
    # 读取 Word 文档
    try:
        doc = Document(input_file)
    except Exception as e:
        print(f"无法读取 {input_file}: {e}")
        return
    
    current_doc = Document()
    chapter_count = 0
    chapter_title = None
    chapter_content = []
    
    for paragraph in doc.paragraphs:
        # 检查是否为指定标题样式
        if paragraph.style.name == heading_style:
            # 如果已经有章节内容，保存当前章节
            if chapter_content or chapter_title:
                save_chapter(current_doc, chapter_title, chapter_content, file_output_dir, chapter_count)
                chapter_count += 1
                current_doc = Document()
                chapter_content = []
            
            # 记录新章节标题
            chapter_title = paragraph.text.strip()
        else:
            # 收集章节内容
            chapter_content.append(paragraph)
    
    # 保存最后一个章节
    if chapter_content or chapter_title:
        save_chapter(current_doc, chapter_title, chapter_content, file_output_dir, chapter_count)
    
    print(f"完成处理 {input_file}，生成 {chapter_count + 1} 个章节文件（样式：{heading_style}）。")

def save_chapter(doc, title, paragraphs, output_dir, chapter_number):
    # 生成安全的文件名
    safe_title = (title.replace('/', '_').replace('\\', '_').replace(':', '_').replace('?', '_').replace('*', '_').replace('\t', '_')
                  if title else f'Chapter_{chapter_number}')
    safe_title = safe_title[:100]  # 限制文件名长度
    docx_path = os.path.join(output_dir, f'{safe_title}.docx')
    
    # 保存 .docx 文件
    if title:
        doc.add_heading(title, level=1)
    for para in paragraphs:
        doc.add_paragraph(para.text, style=para.style)
    doc.save(docx_path)
    print(f'保存 .docx: {docx_path}')

def convert_to_txt(input_file, output_dir, relative_path=''):
    # 创建输出目录
    file_output_dir = os.path.join(output_dir, relative_path)
    if not os.path.exists(file_output_dir):
        os.makedirs(file_output_dir)
    
    # 生成输出文件名
    file_name = os.path.splitext(os.path.basename(input_file))[0]
    txt_path = os.path.join(file_output_dir, f'{file_name}.txt')
    
    # 读取 Word 文档
    try:
        doc = Document(input_file)
    except Exception as e:
        print(f"无法读取 {input_file}: {e}")
        return
    
    # 保存为 .txt 文件
    with open(txt_path, 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # 仅写入非空段落
                f.write(text + '\n\n')
    
    print(f'保存 .txt: {txt_path}')

def process_split_folder(input_folder, output_folder, heading_style):
    # 确保输出目录存在
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # 查找所有 .docx 文件（仅当前文件夹）
    docx_files = glob.glob(os.path.join(input_folder, '*.docx'))
    
    if not docx_files:
        print(f"文件夹 {input_folder} 中未找到 .docx 文件。")
        return
    
    # 处理每个文件
    for docx_file in docx_files:
        print(f"正在处理: {docx_file}")
        split_by_chapters(docx_file, output_folder, heading_style)
    
    print(f"所有文件分割完成（样式：{heading_style}）。")

def process_convert_folder(input_folder, output_folder):
    # 确保输出目录存在
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    # 递归查找所有 .docx 文件
    docx_files = []
    for root, _, files in os.walk(input_folder):
        for file in files:
            if file.endswith('.docx'):
                docx_files.append(os.path.join(root, file))
    
    if not docx_files:
        print(f"文件夹 {input_folder} 及其子文件夹中未找到 .docx 文件。")
        return
    
    # 处理每个文件
    for docx_file in docx_files:
        print(f"正在处理: {docx_file}")
        # 计算相对路径以保持目录结构
        relative_path = os.path.relpath(os.path.dirname(docx_file), input_folder)
        convert_to_txt(docx_file, output_folder, relative_path)
    
    print("所有文件转换完成。")

def main():
    input_directory = "input"  # 替换为你的输入文件夹路径
    output_directory = "output"  # 替换为你的输出文件夹路径
    # 设置命令行参数
    parser = argparse.ArgumentParser(description='处理 Word 文档：按章节分割或转为 TXT')
    parser.add_argument('--mode', required=True, choices=['split', 'convert'],
                        help='操作模式：split（按章节分割）或 convert（转为 TXT）')
    parser.add_argument('--style', type=int, default=1, choices=range(1, 10),
                        help='分割模式下的标题样式（整数，1 表示 Heading 1，2 表示 Heading 2 等），仅对 split 模式有效')
    
    args = parser.parse_args()
    
    # 将整数 heading_level 映射到样式名称（例如 1 -> "Heading 1"）
    heading_style = f'Heading {args.style}'
    # 根据模式执行相应功能
    if args.mode == 'split':
        process_split_folder(input_directory, output_directory, heading_style)
    elif args.mode == 'convert':
        process_convert_folder(input_directory, output_directory)

if __name__ == "__main__":
    main()